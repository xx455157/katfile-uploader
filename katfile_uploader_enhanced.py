#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
KatFile 增強版上傳工具 v3.3
包含檔案壓縮和Word文件記錄功能
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import requests
import json
import os
import threading
import time
from pathlib import Path
from datetime import datetime
import sys
from urllib.parse import urlencode, quote
import socket
import urllib3
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import re
import zipfile
import py7zr
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import shutil

class KatFileUploaderEnhanced:
    def __init__(self, root):
        self.root = root
        self.root.title("KatFile 增強版上傳工具 v3.3")
        self.root.geometry("1200x800")
        
        # 初始化變數
        self.api_key = tk.StringVar()
        self.selected_files = []
        self.folders = []
        self.current_folder_id = 0
        self.account_info = {}
        self.is_uploading = False
        self.upload_records = []  # 上傳記錄
        
        # 壓縮設定
        self.compress_enabled = tk.BooleanVar(value=False)
        self.compress_password = tk.StringVar()
        self.compress_format = tk.StringVar(value="zip")
        
        # Word文件設定
        self.generate_word = tk.BooleanVar(value=True)
        self.word_template_path = ""
        
        # 設定檔路徑
        self.config_file = Path.home() / ".katfile_uploader_config.json"
        
        # 載入設定
        self.load_config()
        
        # 建立GUI
        self.create_widgets()
        
        # 建立改進的請求會話
        self.setup_session()
        
        # 如果有API金鑰，自動載入帳戶資訊
        if self.api_key.get().strip():
            self.load_account_info()
    
    def setup_session(self):
        """設定改進的請求會話"""
        self.session = requests.Session()
        
        # 設定重試策略
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
            allowed_methods=["HEAD", "GET", "OPTIONS", "POST"]
        )
        
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)
        
        # 設定預設標頭
        self.session.headers.update({
            'User-Agent': 'KatFile-Uploader/3.3',
            'Accept': 'application/json'
        })
        
        # 禁用SSL警告
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    def create_widgets(self):
        """建立GUI元件"""
        # 建立筆記本容器
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 主要上傳頁面
        main_frame = ttk.Frame(notebook)
        notebook.add(main_frame, text="📤 檔案上傳")
        
        # 壓縮設定頁面
        compress_frame = ttk.Frame(notebook)
        notebook.add(compress_frame, text="🗜️ 壓縮設定")
        
        # Word文件設定頁面
        word_frame = ttk.Frame(notebook)
        notebook.add(word_frame, text="📄 文件記錄")
        
        # 建立各頁面內容
        self.create_main_page(main_frame)
        self.create_compress_page(compress_frame)
        self.create_word_page(word_frame)
    
    def create_main_page(self, parent):
        """建立主要上傳頁面"""
        # 主要容器
        main_container = ttk.PanedWindow(parent, orient=tk.HORIZONTAL)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # 左側面板
        left_panel = ttk.Frame(main_container)
        main_container.add(left_panel, weight=1)
        
        # 右側面板
        right_panel = ttk.Frame(main_container)
        main_container.add(right_panel, weight=2)
        
        # 建立左側內容
        self.create_left_panel(left_panel)
        
        # 建立右側內容
        self.create_right_panel(right_panel)
    
    def create_compress_page(self, parent):
        """建立壓縮設定頁面"""
        # 壓縮啟用
        enable_frame = ttk.LabelFrame(parent, text="壓縮設定", padding="10")
        enable_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Checkbutton(
            enable_frame, 
            text="啟用檔案壓縮（上傳前自動壓縮檔案）", 
            variable=self.compress_enabled
        ).pack(anchor=tk.W)
        
        # 壓縮格式
        format_frame = ttk.LabelFrame(parent, text="壓縮格式", padding="10")
        format_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Radiobutton(format_frame, text="ZIP格式（相容性好）", variable=self.compress_format, value="zip").pack(anchor=tk.W)
        ttk.Radiobutton(format_frame, text="7Z格式（壓縮率高）", variable=self.compress_format, value="7z").pack(anchor=tk.W)
        
        # 密碼設定
        password_frame = ttk.LabelFrame(parent, text="壓縮密碼", padding="10")
        password_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(password_frame, text="壓縮密碼（留空則不加密）:").pack(anchor=tk.W)
        
        password_entry_frame = ttk.Frame(password_frame)
        password_entry_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.password_entry = ttk.Entry(password_entry_frame, textvariable=self.compress_password, show="*", width=30)
        self.password_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.show_password_var = tk.BooleanVar()
        ttk.Checkbutton(
            password_entry_frame, 
            text="顯示", 
            variable=self.show_password_var,
            command=self.toggle_password_visibility
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        # 測試壓縮
        test_frame = ttk.LabelFrame(parent, text="測試壓縮", padding="10")
        test_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(test_frame, text="🧪 測試壓縮功能", command=self.test_compression).pack()
        
        # 壓縮說明
        info_frame = ttk.LabelFrame(parent, text="說明", padding="10")
        info_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        info_text = """壓縮功能說明：

1. 啟用壓縮後，上傳前會自動將檔案壓縮
2. ZIP格式相容性好，支援所有系統
3. 7Z格式壓縮率更高，但需要專用軟體解壓
4. 設定密碼後，解壓時需要輸入密碼
5. 壓縮後的檔案會自動上傳，原檔案保持不變
6. 大檔案壓縮可能需要較長時間，請耐心等待"""
        
        info_label = tk.Text(info_frame, wrap=tk.WORD, height=10, state=tk.DISABLED)
        info_label.pack(fill=tk.BOTH, expand=True)
        info_label.config(state=tk.NORMAL)
        info_label.insert(tk.END, info_text)
        info_label.config(state=tk.DISABLED)
    
    def create_word_page(self, parent):
        """建立Word文件設定頁面"""
        # Word文件啟用
        enable_frame = ttk.LabelFrame(parent, text="Word文件記錄", padding="10")
        enable_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Checkbutton(
            enable_frame, 
            text="自動生成Word文件記錄（每個檔案生成一份記錄）", 
            variable=self.generate_word
        ).pack(anchor=tk.W)
        
        # 範本設定
        template_frame = ttk.LabelFrame(parent, text="範本設定", padding="10")
        template_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(template_frame, text="Word範本檔案:").pack(anchor=tk.W)
        
        template_path_frame = ttk.Frame(template_frame)
        template_path_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.template_path_var = tk.StringVar()
        ttk.Entry(template_path_frame, textvariable=self.template_path_var, state="readonly").pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(template_path_frame, text="瀏覽", command=self.select_word_template).pack(side=tk.RIGHT, padx=(5, 0))
        
        # 預設範本
        default_frame = ttk.Frame(template_frame)
        default_frame.pack(fill=tk.X, pady=(10, 0))
        
        ttk.Button(default_frame, text="📄 使用內建範本", command=self.use_builtin_template).pack(side=tk.LEFT)
        ttk.Button(default_frame, text="🔍 預覽範本", command=self.preview_template).pack(side=tk.LEFT, padx=(10, 0))
        
        # 輸出設定
        output_frame = ttk.LabelFrame(parent, text="輸出設定", padding="10")
        output_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(output_frame, text="Word文件將儲存到與原檔案相同的目錄").pack(anchor=tk.W)
        ttk.Label(output_frame, text="檔案名稱格式：[原檔案名]_記錄.docx").pack(anchor=tk.W)
        
        # Word文件說明
        word_info_frame = ttk.LabelFrame(parent, text="說明", padding="10")
        word_info_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        word_info_text = """Word文件記錄功能說明：

1. 每個上傳的檔案都會生成一份Word記錄文件
2. 記錄包含：檔案名稱、格式、大小、下載連結等資訊
3. 如果啟用壓縮，會自動填入解壓密碼
4. 可以使用自訂範本或內建範本
5. Word文件會儲存在與原檔案相同的目錄
6. 支援批次生成，一次上傳多個檔案會生成多份記錄

內建範本包含以下欄位：
- 檔案名稱
- 影片格式
- 影片大小  
- 影片說明
- 解壓密碼
- 影片載點（含超連結）
- 影片截圖區域"""
        
        word_info_label = tk.Text(word_info_frame, wrap=tk.WORD, height=12, state=tk.DISABLED)
        word_info_label.pack(fill=tk.BOTH, expand=True)
        word_info_label.config(state=tk.NORMAL)
        word_info_label.insert(tk.END, word_info_text)
        word_info_label.config(state=tk.DISABLED)
    
    def toggle_password_visibility(self):
        """切換密碼顯示/隱藏"""
        if self.show_password_var.get():
            self.password_entry.config(show="")
        else:
            self.password_entry.config(show="*")
    
    def test_compression(self):
        """測試壓縮功能"""
        if not self.compress_enabled.get():
            messagebox.showinfo("提示", "請先啟用壓縮功能")
            return
        
        # 選擇測試檔案
        test_file = filedialog.askopenfilename(
            title="選擇測試檔案",
            filetypes=[("所有檔案", "*.*")]
        )
        
        if not test_file:
            return
        
        def test_thread():
            try:
                self.log("🧪 開始測試壓縮...")
                
                # 建立測試輸出目錄
                test_dir = Path.home() / "katfile_compression_test"
                test_dir.mkdir(exist_ok=True)
                
                # 壓縮檔案
                compressed_file = self.compress_file(test_file, test_dir)
                
                if compressed_file:
                    success_msg = f"✅ 壓縮測試成功！\n壓縮檔案：{compressed_file}"
                    self.root.after(0, lambda: messagebox.showinfo("測試成功", success_msg))
                    self.root.after(0, lambda: self.log(f"✅ 測試壓縮成功：{compressed_file}"))
                else:
                    self.root.after(0, lambda: messagebox.showerror("測試失敗", "壓縮測試失敗"))
                    
            except Exception as e:
                error_msg = f"❌ 壓縮測試錯誤：{str(e)}"
                self.root.after(0, lambda: self.log(error_msg))
                self.root.after(0, lambda: messagebox.showerror("錯誤", error_msg))
        
        threading.Thread(target=test_thread, daemon=True).start()
    
    def select_word_template(self):
        """選擇Word範本檔案"""
        template_file = filedialog.askopenfilename(
            title="選擇Word範本檔案",
            filetypes=[("Word文件", "*.docx"), ("所有檔案", "*.*")]
        )
        
        if template_file:
            self.template_path_var.set(template_file)
            self.word_template_path = template_file
            self.log(f"📄 已選擇Word範本：{template_file}")
    
    def use_builtin_template(self):
        """使用內建範本"""
        self.template_path_var.set("內建範本")
        self.word_template_path = ""
        self.log("📄 已選擇內建Word範本")
    
    def preview_template(self):
        """預覽範本"""
        if self.word_template_path and os.path.exists(self.word_template_path):
            # 開啟自訂範本
            try:
                os.startfile(self.word_template_path)  # Windows
            except:
                try:
                    os.system(f'open "{self.word_template_path}"')  # macOS
                except:
                    os.system(f'xdg-open "{self.word_template_path}"')  # Linux
        else:
            # 顯示內建範本說明
            preview_text = """內建範本格式：

{檔案名稱}@MP4@KF@無碼

【影片名稱】：{檔案名稱}
【影片格式】：MP4
【影片大小】：{影片大小}
【影片說明】：無碼
【解壓密碼】：{解壓縮密碼}
【影片載點】：{檔案名稱+網址的超連結}
【影片截圖】：

我的伊利所有帖子

破處, 國產, 學妹, 蘿莉, 處女"""
            
            messagebox.showinfo("內建範本預覽", preview_text)
    
    def create_left_panel(self, parent):
        """建立左側面板（API設定和帳戶資訊）"""
        # API金鑰設定區域
        api_frame = ttk.LabelFrame(parent, text="API金鑰設定", padding="10")
        api_frame.pack(fill=tk.X, pady=(0, 10))
        
        # API金鑰輸入
        ttk.Label(api_frame, text="API金鑰:").pack(anchor=tk.W)
        
        key_frame = ttk.Frame(api_frame)
        key_frame.pack(fill=tk.X, pady=(5, 10))
        
        self.api_entry = ttk.Entry(key_frame, textvariable=self.api_key, show="*", width=30)
        self.api_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 顯示/隱藏按鈕
        self.show_key_var = tk.BooleanVar()
        self.show_key_btn = ttk.Checkbutton(
            key_frame, 
            text="顯示", 
            variable=self.show_key_var,
            command=self.toggle_key_visibility
        )
        self.show_key_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # 按鈕區域
        api_buttons = ttk.Frame(api_frame)
        api_buttons.pack(fill=tk.X)
        
        ttk.Button(api_buttons, text="儲存", command=self.save_api_key).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(api_buttons, text="測試", command=self.test_api_key).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(api_buttons, text="診斷", command=self.diagnose_network).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(api_buttons, text="清除", command=self.clear_api_key).pack(side=tk.LEFT)
        
        # 帳戶資訊區域
        account_frame = ttk.LabelFrame(parent, text="帳戶資訊", padding="10")
        account_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        self.account_text = scrolledtext.ScrolledText(account_frame, height=8, state=tk.DISABLED)
        self.account_text.pack(fill=tk.BOTH, expand=True)
        
        # 資料夾管理區域
        folder_frame = ttk.LabelFrame(parent, text="資料夾管理", padding="10")
        folder_frame.pack(fill=tk.BOTH, expand=True)
        
        # 資料夾操作按鈕
        folder_buttons = ttk.Frame(folder_frame)
        folder_buttons.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(folder_buttons, text="🔄", command=self.refresh_folders, width=3).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(folder_buttons, text="➕", command=self.create_folder, width=3).pack(side=tk.LEFT)
        
        # 資料夾列表
        self.folder_tree = ttk.Treeview(folder_frame, columns=("id",), show="tree")
        self.folder_tree.pack(fill=tk.BOTH, expand=True)
        self.folder_tree.bind("<Double-1>", self.on_folder_select)
    
    def create_right_panel(self, parent):
        """建立右側面板（檔案上傳和日誌）"""
        # 檔案選擇區域
        file_frame = ttk.LabelFrame(parent, text="檔案上傳", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 目標資料夾選擇
        target_frame = ttk.Frame(file_frame)
        target_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(target_frame, text="上傳到:").pack(side=tk.LEFT)
        self.target_folder_var = tk.StringVar(value="根目錄")
        self.target_folder_label = ttk.Label(target_frame, textvariable=self.target_folder_var, foreground="blue")
        self.target_folder_label.pack(side=tk.LEFT, padx=(10, 0))
        
        # 檔案選擇按鈕
        button_frame = ttk.Frame(file_frame)
        button_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(button_frame, text="📁 選擇檔案", command=self.select_files).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="📂 選擇資料夾", command=self.select_folder).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="🗑️ 清除列表", command=self.clear_files).pack(side=tk.LEFT)
        
        # 檔案列表
        self.file_tree = ttk.Treeview(file_frame, columns=("size", "status"), show="tree headings", height=6)
        self.file_tree.heading("#0", text="檔案名稱")
        self.file_tree.heading("size", text="大小")
        self.file_tree.heading("status", text="狀態")
        self.file_tree.column("size", width=100)
        self.file_tree.column("status", width=100)
        self.file_tree.pack(fill=tk.X)
        
        # 上傳控制
        upload_frame = ttk.Frame(file_frame)
        upload_frame.pack(fill=tk.X, pady=(10, 0))
        
        self.upload_button = ttk.Button(upload_frame, text="🚀 開始上傳", command=self.start_upload)
        self.upload_button.pack(side=tk.LEFT)
        
        self.progress = ttk.Progressbar(upload_frame, mode='determinate')
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(10, 0))
        
        # 日誌區域
        log_frame = ttk.LabelFrame(parent, text="操作日誌", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=15)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # 日誌控制按鈕
        log_buttons = ttk.Frame(log_frame)
        log_buttons.pack(fill=tk.X, pady=(10, 0))
        
        ttk.Button(log_buttons, text="清除日誌", command=self.clear_log).pack(side=tk.LEFT)
        ttk.Button(log_buttons, text="儲存日誌", command=self.save_log).pack(side=tk.LEFT, padx=(10, 0))
        ttk.Button(log_buttons, text="📄 生成Word報告", command=self.generate_upload_report).pack(side=tk.LEFT, padx=(10, 0))
    
    def add_hyperlink(self, paragraph, url, text):
        """在段落中添加超連結"""
        try:
            # 建立超連結元素
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
            
            # 建立文字運行
            run = OxmlElement('w:r')
            
            # 設定超連結樣式
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # 藍色
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(color)
            rPr.append(underline)
            run.append(rPr)
            
            # 添加文字
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            run.append(text_elem)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            
            return True
        except Exception as e:
            print(f"建立超連結失敗: {e}")
            return False

    def compress_file(self, file_path, output_dir):
        """壓縮檔案"""
        try:
            file_path = Path(file_path)
            output_dir = Path(output_dir)
            
            # 建立壓縮檔案名稱
            if self.compress_format.get() == "zip":
                compressed_file = output_dir / f"{file_path.stem}.zip"
            else:
                compressed_file = output_dir / f"{file_path.stem}.7z"
            
            password = self.compress_password.get().strip() if self.compress_password.get().strip() else None
            
            self.log(f"🗜️ 開始壓縮：{file_path.name}")
            
            if self.compress_format.get() == "zip":
                # ZIP壓縮
                with zipfile.ZipFile(compressed_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    if password:
                        zipf.setpassword(password.encode('utf-8'))
                    zipf.write(file_path, file_path.name)
            else:
                # 7Z壓縮
                with py7zr.SevenZipFile(compressed_file, 'w', password=password) as archive:
                    archive.write(file_path, file_path.name)
            
            self.log(f"✅ 壓縮完成：{compressed_file.name}")
            return str(compressed_file)
            
        except Exception as e:
            self.log(f"❌ 壓縮失敗：{str(e)}")
            return None
    
    def generate_word_document(self, file_info, download_link, compressed_file=None):
        """生成Word文件記錄"""
        try:
            if self.word_template_path and os.path.exists(self.word_template_path):
                # 使用自訂範本
                doc = Document(self.word_template_path)
            else:
                # 使用內建範本
                doc = Document()
                
                # 建立標題
                title = doc.add_heading(f"{file_info['name']}@MP4@KF@無碼", level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加空行
                doc.add_paragraph()
                
                # 建立資訊表格
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'
                
                # 填入資訊
                cells = table.rows[0].cells
                cells[0].text = "【影片名稱】"
                cells[1].text = f"：{file_info['name']}"
                
                cells = table.rows[1].cells
                cells[0].text = "【影片格式】"
                cells[1].text = "：MP4"
                
                cells = table.rows[2].cells
                cells[0].text = "【影片大小】"
                cells[1].text = f"：{self.format_file_size(file_info['size'])}"
                
                cells = table.rows[3].cells
                cells[0].text = "【影片說明】"
                cells[1].text = "：無碼"
                
                cells = table.rows[4].cells
                cells[0].text = "【解壓密碼】"
                if compressed_file and self.compress_password.get().strip():
                    cells[1].text = f"：{self.compress_password.get().strip()}"
                else:
                    cells[1].text = "：無"
                
                cells = table.rows[5].cells
                cells[0].text = "【影片載點】"
                # 添加超連結
                paragraph = cells[1].paragraphs[0]
                paragraph.text = "："
                
                # 建立超連結
                hyperlink = self.add_hyperlink(paragraph, download_link, file_info['name'])
                if not hyperlink:
                    # 如果超連結建立失敗，至少顯示連結文字
                    run = paragraph.add_run(f"{file_info['name']} - {download_link}")
                    run.font.color.rgb = RGBColor(0, 0, 255)  # 藍色文字
                
                # 添加空行和截圖區域
                doc.add_paragraph()
                doc.add_paragraph("【影片截圖】：")
                doc.add_paragraph()
                
                # 添加底部標題
                footer_title = doc.add_heading("我的伊利所有帖子", level=2)
                footer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加標籤
                tags_para = doc.add_paragraph("破處, 國產, 學妹, 蘿莉, 處女")
                tags_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 儲存Word文件
            file_dir = Path(file_info['path']).parent
            word_filename = f"{Path(file_info['name']).stem}_記錄.docx"
            word_path = file_dir / word_filename
            
            doc.save(str(word_path))
            
            self.log(f"📄 Word文件已生成：{word_path}")
            return str(word_path)
            
        except Exception as e:
            self.log(f"❌ Word文件生成失敗：{str(e)}")
            return None
    
    def generate_upload_report(self):
        """生成上傳報告"""
        if not self.upload_records:
            messagebox.showinfo("提示", "沒有上傳記錄可生成報告")
            return
        
        try:
            # 選擇儲存位置
            report_file = filedialog.asksaveasfilename(
                title="儲存上傳報告",
                defaultextension=".docx",
                filetypes=[("Word文件", "*.docx"), ("所有檔案", "*.*")]
            )
            
            if not report_file:
                return
            
            # 建立報告文件
            doc = Document()
            
            # 標題
            title = doc.add_heading("KatFile 上傳報告", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 生成時間
            doc.add_paragraph(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"總計上傳檔案：{len(self.upload_records)} 個")
            
            # 建立表格
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # 表頭
            header_cells = table.rows[0].cells
            header_cells[0].text = "檔案名稱"
            header_cells[1].text = "檔案大小"
            header_cells[2].text = "上傳時間"
            header_cells[3].text = "下載連結"
            header_cells[4].text = "狀態"
            
            # 填入記錄
            for record in self.upload_records:
                row_cells = table.add_row().cells
                row_cells[0].text = record.get('filename', 'N/A')
                row_cells[1].text = record.get('filesize', 'N/A')
                row_cells[2].text = record.get('upload_time', 'N/A')
                row_cells[3].text = record.get('download_link', 'N/A')
                row_cells[4].text = record.get('status', 'N/A')
            
            doc.save(report_file)
            
            self.log(f"📊 上傳報告已生成：{report_file}")
            messagebox.showinfo("成功", f"上傳報告已儲存到：{report_file}")
            
        except Exception as e:
            error_msg = f"❌ 生成報告失敗：{str(e)}"
            self.log(error_msg)
            messagebox.showerror("錯誤", error_msg)
    
    # 以下是原有的方法，保持不變
    def toggle_key_visibility(self):
        """切換API金鑰顯示/隱藏"""
        if self.show_key_var.get():
            self.api_entry.config(show="")
        else:
            self.api_entry.config(show="*")
    
    def clear_api_key(self):
        """清除API金鑰"""
        if messagebox.askyesno("確認", "確定要清除API金鑰嗎？"):
            self.api_key.set("")
            self.save_config()
            self.account_text.config(state=tk.NORMAL)
            self.account_text.delete(1.0, tk.END)
            self.account_text.config(state=tk.DISABLED)
            self.log("🗑️ API金鑰已清除")
    
    def load_config(self):
        """載入設定"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.api_key.set(config.get('api_key', ''))
                    self.compress_enabled.set(config.get('compress_enabled', False))
                    self.compress_password.set(config.get('compress_password', ''))
                    self.compress_format.set(config.get('compress_format', 'zip'))
                    self.generate_word.set(config.get('generate_word', True))
                    self.word_template_path = config.get('word_template_path', '')
        except Exception as e:
            self.log(f"⚠️ 載入設定失敗: {e}")
    
    def save_config(self):
        """儲存設定"""
        try:
            config = {
                'api_key': self.api_key.get().strip(),
                'compress_enabled': self.compress_enabled.get(),
                'compress_password': self.compress_password.get(),
                'compress_format': self.compress_format.get(),
                'generate_word': self.generate_word.get(),
                'word_template_path': self.word_template_path
            }
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.log(f"⚠️ 儲存設定失敗: {e}")
    
    def log(self, message):
        """記錄日誌"""
        timestamp = datetime.now().strftime("[%H:%M:%S]")
        log_message = f"{timestamp} {message}\n"
        
        self.log_text.insert(tk.END, log_message)
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def clear_log(self):
        """清除日誌"""
        self.log_text.delete(1.0, tk.END)
    
    def save_log(self):
        """儲存日誌"""
        try:
            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("文字檔案", "*.txt"), ("所有檔案", "*.*")]
            )
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(self.log_text.get(1.0, tk.END))
                self.log(f"📄 日誌已儲存到: {filename}")
        except Exception as e:
            self.log(f"❌ 儲存日誌失敗: {e}")
    
    def validate_api_key(self, key):
        """驗證API金鑰格式"""
        if not key:
            return False, "API金鑰不能為空"
        
        key = key.strip()
        
        if len(key) < 10:
            return False, "API金鑰長度太短"
        
        if not re.match(r'^[a-zA-Z0-9]+$', key):
            return False, "API金鑰包含無效字符，只能包含字母和數字"
        
        return True, "格式正確"
    
    def save_api_key(self):
        """儲存API金鑰"""
        key = self.api_key.get().strip()
        
        if not key:
            messagebox.showwarning("警告", "請先輸入API金鑰！")
            return
        
        is_valid, message = self.validate_api_key(key)
        if not is_valid:
            messagebox.showerror("錯誤", f"API金鑰格式錯誤: {message}")
            return
        
        self.api_key.set(key)
        self.save_config()
        self.log("✅ API金鑰已儲存")
        
        self.test_api_key()
    
    def test_api_key(self):
        """測試API金鑰"""
        key = self.api_key.get().strip()
        
        if not key:
            messagebox.showwarning("警告", "請先輸入API金鑰！")
            return
        
        is_valid, message = self.validate_api_key(key)
        if not is_valid:
            messagebox.showerror("錯誤", f"API金鑰格式錯誤: {message}")
            return
            
        self.log("🔍 測試API金鑰...")
        
        def test_thread():
            try:
                url = f"https://katfile.cloud/api/account/info?key={quote(key)}"
                response = self.session.get(url, timeout=15, allow_redirects=True)
                
                if response.status_code == 200:
                    data = response.json()
                    if data.get('msg') == 'OK':
                        self.root.after(0, lambda: self.log("✅ API金鑰測試成功"))
                        self.root.after(0, lambda: messagebox.showinfo("成功", "API金鑰有效！"))
                        self.root.after(0, self.load_account_info)
                        return
                    else:
                        error_msg = f"❌ API金鑰無效: {data.get('msg', '未知錯誤')}"
                        self.root.after(0, lambda msg=error_msg: self.log(msg))
                        self.root.after(0, lambda: messagebox.showerror("錯誤", f"API金鑰無效: {data.get('msg', '未知錯誤')}"))
                        return
                        
            except Exception as error:
                error_msg = f"❌ 測試失敗: {str(error)}"
                self.root.after(0, lambda msg=error_msg: self.log(msg))
                self.root.after(0, lambda: messagebox.showerror("錯誤", "測試失敗，請檢查API金鑰和網路連線"))
                
        threading.Thread(target=test_thread, daemon=True).start()
    
    def diagnose_network(self):
        """診斷網路連線"""
        self.log("🔍 開始網路診斷...")
        
        def diagnose_thread():
            try:
                ip = socket.gethostbyname('katfile.cloud')
                self.root.after(0, lambda: self.log(f"✅ DNS解析成功: katfile.cloud -> {ip}"))
                
                response = self.session.get('https://katfile.cloud', timeout=10, allow_redirects=True)
                self.root.after(0, lambda: self.log(f"✅ 基本連線成功: HTTP {response.status_code}"))
                
                response = self.session.get("https://katfile.cloud/api/account/info?key=test", timeout=10, allow_redirects=True)
                if response.status_code in [200, 400, 401]:
                    self.root.after(0, lambda: self.log("✅ API端點可正常訪問"))
                else:
                    self.root.after(0, lambda: self.log(f"⚠️ API端點回應異常: HTTP {response.status_code}"))
                    
            except Exception as e:
                self.root.after(0, lambda: self.log(f"❌ 網路診斷失敗: {str(e)}"))
        
        threading.Thread(target=diagnose_thread, daemon=True).start()
    
    def load_account_info(self):
        """載入帳戶資訊"""
        key = self.api_key.get().strip()
        if not key:
            return
            
        def load_thread():
            try:
                url = f"https://katfile.cloud/api/account/info?key={quote(key)}"
                response = self.session.get(url, timeout=15, allow_redirects=True)
                
                if response.status_code == 200:
                    self.account_info = response.json()
                    self.root.after(0, self.display_account_info)
                    
                self.root.after(0, self.refresh_folders)
                    
            except Exception as error:
                error_msg = f"❌ 載入帳戶資訊失敗: {str(error)}"
                self.root.after(0, lambda msg=error_msg: self.log(msg))
                
        threading.Thread(target=load_thread, daemon=True).start()
    
    def display_account_info(self):
        """顯示帳戶資訊"""
        if not self.account_info:
            return
            
        self.account_text.config(state=tk.NORMAL)
        self.account_text.delete(1.0, tk.END)
        
        result = self.account_info.get('result', {})
        
        storage_used = int(result.get('storage_used', 0))
        storage_left = int(result.get('storage_left', 0))
        storage_total = storage_used + storage_left
        
        def format_bytes(bytes_val):
            for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
                if bytes_val < 1024.0:
                    return f"{bytes_val:.2f} {unit}"
                bytes_val /= 1024.0
            return f"{bytes_val:.2f} PB"
        
        account_display = f"""📧 Email: {result.get('email', 'N/A')}
💰 餘額: ${result.get('balance', 'N/A')}
⭐ Premium到期: {result.get('premium_expire', 'N/A')}
💾 已使用空間: {format_bytes(storage_used)}
💿 剩餘空間: {format_bytes(storage_left)}
📊 總空間: {format_bytes(storage_total)}
🕒 伺服器時間: {self.account_info.get('server_time', 'N/A')}"""
        
        self.account_text.insert(tk.END, account_display)
        self.account_text.config(state=tk.DISABLED)
    
    def refresh_folders(self):
        """重新整理資料夾列表"""
        key = self.api_key.get().strip()
        if not key:
            return
            
        def refresh_thread():
            try:
                url = f"https://katfile.cloud/api/folder/list?key={quote(key)}"
                response = self.session.get(url, timeout=15, allow_redirects=True)
                
                if response.status_code == 200:
                    data = response.json()
                    if data.get('msg') == 'OK':
                        self.folders = data.get('result', {}).get('folders', [])
                        self.root.after(0, self.update_folder_display)
                        
            except Exception as error:
                error_msg = f"❌ 載入資料夾失敗: {str(error)}"
                self.root.after(0, lambda msg=error_msg: self.log(msg))
                
        threading.Thread(target=refresh_thread, daemon=True).start()
    
    def update_folder_display(self):
        """更新資料夾顯示"""
        for item in self.folder_tree.get_children():
            self.folder_tree.delete(item)
        
        root_item = self.folder_tree.insert("", "end", text="📁 根目錄", values=(0,))
        
        for folder in self.folders:
            folder_name = folder.get('name', '未知資料夾')
            folder_id = folder.get('fld_id', 0)
            self.folder_tree.insert("", "end", text=f"📁 {folder_name}", values=(folder_id,))
        
        self.log(f"📁 載入了 {len(self.folders)} 個資料夾")
    
    def on_folder_select(self, event):
        """資料夾選擇事件"""
        selection = self.folder_tree.selection()
        if selection:
            item = selection[0]
            folder_id = self.folder_tree.item(item, "values")[0]
            folder_name = self.folder_tree.item(item, "text")
            
            self.current_folder_id = int(folder_id)
            self.target_folder_var.set(folder_name.replace("📁 ", ""))
            
            self.log(f"📂 選擇目標資料夾: {folder_name} (ID: {folder_id})")
    
    def create_folder(self):
        """建立新資料夾"""
        key = self.api_key.get().strip()
        if not key:
            messagebox.showwarning("警告", "請先設定API金鑰！")
            return
        
        folder_name = simpledialog.askstring("建立資料夾", "請輸入資料夾名稱:")
        if not folder_name:
            return
            
        def create_thread():
            try:
                params = {
                    'key': key,
                    'name': folder_name
                }
                
                if self.current_folder_id != 0:
                    params['parent_id'] = self.current_folder_id
                
                url = f"https://katfile.cloud/api/folder/create?{urlencode(params)}"
                response = self.session.get(url, timeout=15, allow_redirects=True)
                
                if response.status_code == 200:
                    data = response.json()
                    if data.get('msg') == 'OK':
                        success_msg = f"✅ 資料夾 '{folder_name}' 建立成功"
                        self.root.after(0, lambda msg=success_msg: self.log(msg))
                        self.root.after(0, self.refresh_folders)
                    else:
                        error_msg = f"❌ 建立資料夾失敗: {data.get('msg', '未知錯誤')}"
                        self.root.after(0, lambda msg=error_msg: self.log(msg))
                        
            except Exception as error:
                error_msg = f"❌ 建立資料夾錯誤: {str(error)}"
                self.root.after(0, lambda msg=error_msg: self.log(msg))
                
        threading.Thread(target=create_thread, daemon=True).start()
    
    def select_files(self):
        """選擇檔案"""
        files = filedialog.askopenfilenames(
            title="選擇要上傳的檔案",
            filetypes=[("所有檔案", "*.*")]
        )
        
        for file_path in files:
            if file_path not in [f['path'] for f in self.selected_files]:
                file_info = {
                    'path': file_path,
                    'name': os.path.basename(file_path),
                    'size': os.path.getsize(file_path)
                }
                self.selected_files.append(file_info)
        
        self.update_file_display()
    
    def select_folder(self):
        """選擇資料夾"""
        folder_path = filedialog.askdirectory(title="選擇要上傳的資料夾")
        
        if folder_path:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if file_path not in [f['path'] for f in self.selected_files]:
                        file_info = {
                            'path': file_path,
                            'name': os.path.relpath(file_path, folder_path),
                            'size': os.path.getsize(file_path)
                        }
                        self.selected_files.append(file_info)
            
            self.update_file_display()
    
    def clear_files(self):
        """清除檔案列表"""
        self.selected_files = []
        self.update_file_display()
        self.log("🗑️ 檔案列表已清除")
    
    def update_file_display(self):
        """更新檔案顯示"""
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        
        for i, file_info in enumerate(self.selected_files):
            size_str = self.format_file_size(file_info['size'])
            self.file_tree.insert("", "end", text=file_info['name'], 
                                values=(size_str, "等待上傳"), tags=(str(i),))
        
        self.log(f"📄 選擇了 {len(self.selected_files)} 個檔案")
    
    def format_file_size(self, size):
        """格式化檔案大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"
    
    def start_upload(self):
        """開始上傳"""
        if not self.selected_files:
            messagebox.showwarning("警告", "請先選擇要上傳的檔案！")
            return
        
        key = self.api_key.get().strip()
        if not key:
            messagebox.showwarning("警告", "請先設定API金鑰！")
            return
        
        if self.is_uploading:
            messagebox.showinfo("提示", "正在上傳中，請稍候...")
            return
        
        # 儲存設定
        self.save_config()
        
        self.is_uploading = True
        self.upload_button.config(text="⏸️ 上傳中...", state='disabled')
        self.progress['maximum'] = len(self.selected_files)
        self.progress['value'] = 0
        
        # 清除上傳記錄
        self.upload_records = []
        
        target_folder_name = self.target_folder_var.get()
        self.log(f"🚀 開始上傳 {len(self.selected_files)} 個檔案到 {target_folder_name}")
        
        if self.compress_enabled.get():
            self.log("🗜️ 壓縮功能已啟用")
        
        if self.generate_word.get():
            self.log("📄 Word文件記錄功能已啟用")
        
        def upload_thread():
            try:
                success_count = 0
                temp_dir = Path.home() / "katfile_temp_compress"
                temp_dir.mkdir(exist_ok=True)
                
                for i, file_info in enumerate(self.selected_files):
                    if not self.is_uploading:
                        break
                    
                    self.root.after(0, lambda idx=i: self.update_file_status(idx, "處理中..."))
                    
                    # 檔案處理（壓縮）
                    upload_file_path = file_info['path']
                    compressed_file = None
                    
                    if self.compress_enabled.get():
                        self.root.after(0, lambda idx=i: self.update_file_status(idx, "壓縮中..."))
                        compressed_file = self.compress_file(file_info['path'], temp_dir)
                        if compressed_file:
                            upload_file_path = compressed_file
                            # 更新檔案資訊用於上傳
                            upload_file_info = {
                                'path': upload_file_path,
                                'name': os.path.basename(upload_file_path),
                                'size': os.path.getsize(upload_file_path)
                            }
                        else:
                            self.root.after(0, lambda idx=i: self.update_file_status(idx, "❌ 壓縮失敗"))
                            continue
                    else:
                        upload_file_info = file_info
                    
                    # 上傳檔案
                    self.root.after(0, lambda idx=i: self.update_file_status(idx, "上傳中..."))
                    
                    download_link = self.upload_single_file(upload_file_info, self.current_folder_id)
                    
                    if download_link:
                        success_count += 1
                        self.root.after(0, lambda idx=i: self.update_file_status(idx, "✅ 完成"))
                        
                        # 記錄上傳資訊
                        upload_record = {
                            'filename': file_info['name'],
                            'filesize': self.format_file_size(file_info['size']),
                            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            'download_link': download_link,
                            'status': '成功'
                        }
                        self.upload_records.append(upload_record)
                        
                        success_msg = f"✅ 上傳成功: {file_info['name']}"
                        self.root.after(0, lambda msg=success_msg: self.log(msg))
                        link_msg = f"🔗 下載連結: {download_link}"
                        self.root.after(0, lambda msg=link_msg: self.log(msg))
                        
                        # 生成Word文件
                        if self.generate_word.get():
                            self.root.after(0, lambda idx=i: self.update_file_status(idx, "生成文件..."))
                            word_file = self.generate_word_document(file_info, download_link, compressed_file)
                            if word_file:
                                word_msg = f"📄 Word文件: {word_file}"
                                self.root.after(0, lambda msg=word_msg: self.log(msg))
                    else:
                        self.root.after(0, lambda idx=i: self.update_file_status(idx, "❌ 失敗"))
                        
                        # 記錄失敗資訊
                        upload_record = {
                            'filename': file_info['name'],
                            'filesize': self.format_file_size(file_info['size']),
                            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            'download_link': 'N/A',
                            'status': '失敗'
                        }
                        self.upload_records.append(upload_record)
                    
                    # 清理臨時壓縮檔案
                    if compressed_file and os.path.exists(compressed_file):
                        try:
                            os.remove(compressed_file)
                        except:
                            pass
                    
                    self.root.after(0, lambda: self.progress.step())
                
                # 清理臨時目錄
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
                
                # 上傳完成
                completion_msg = f"🎉 上傳完成！成功: {success_count}/{len(self.selected_files)}"
                self.root.after(0, lambda msg=completion_msg: self.log(msg))
                
            finally:
                self.is_uploading = False
                self.root.after(0, lambda: self.upload_button.config(text="🚀 開始上傳", state='normal'))
        
        threading.Thread(target=upload_thread, daemon=True).start()
    
    def upload_single_file(self, file_info, target_folder_id):
        """上傳單個檔案"""
        key = self.api_key.get().strip()
        max_retries = 2
        
        for attempt in range(max_retries):
            try:
                if attempt > 0:
                    retry_msg = f"🔄 重試上傳 {file_info['name']} (第 {attempt} 次)"
                    self.root.after(0, lambda msg=retry_msg: self.log(msg))
                    time.sleep(3)
                
                # 第一步：獲取上傳伺服器
                server_url = f"https://katfile.cloud/api/upload/server?key={quote(key)}"
                response = self.session.get(server_url, timeout=30, allow_redirects=True)
                
                if response.status_code != 200:
                    raise Exception(f"獲取上傳伺服器失敗: HTTP {response.status_code}")
                    
                upload_context = response.json()
                if upload_context.get('msg') != 'OK':
                    raise Exception(f"API錯誤: {upload_context.get('msg', '未知錯誤')}")
                    
                # 第二步：上傳檔案
                upload_url = upload_context['result']
                sess_id = upload_context['sess_id']
                
                with open(file_info['path'], 'rb') as f:
                    files = {'file_0': (file_info['name'], f, 'application/octet-stream')}
                    data = {
                        'sess_id': sess_id,
                        'utype': 'prem'
                    }
                    
                    response = self.session.post(
                        upload_url, 
                        files=files, 
                        data=data, 
                        timeout=600,
                        allow_redirects=True
                    )
                    
                if response.status_code != 200:
                    raise Exception(f"上傳失敗: HTTP {response.status_code}")
                    
                upload_result = response.json()
                if not upload_result or not isinstance(upload_result, list):
                    raise Exception("上傳回應格式錯誤")
                    
                file_result = upload_result[0]
                if file_result.get('file_status') != 'OK':
                    raise Exception(f"上傳失敗: {file_result.get('file_status', '未知錯誤')}")
                    
                file_code = file_result['file_code']
                
                # 第三步：移動到目標資料夾
                if target_folder_id != 0:
                    try:
                        move_url = f"https://katfile.cloud/api/file/set_folder?key={quote(key)}&file_code={file_code}&fld_id={target_folder_id}"
                        move_response = self.session.get(move_url, timeout=30, allow_redirects=True)
                        
                        if move_response.status_code != 200:
                            warning_msg = f"⚠️ 移動檔案到資料夾失敗: HTTP {move_response.status_code}"
                            self.root.after(0, lambda msg=warning_msg: self.log(msg))
                    except:
                        warning_msg = f"⚠️ 移動檔案到資料夾時發生錯誤，檔案已上傳到根目錄"
                        self.root.after(0, lambda msg=warning_msg: self.log(msg))
                        
                # 第四步：獲取直接下載連結
                try:
                    direct_url = f"https://katfile.cloud/api/file/direct_link?key={quote(key)}&file_code={file_code}"
                    direct_response = self.session.get(direct_url, timeout=30, allow_redirects=True)
                    
                    if direct_response.status_code == 200:
                        direct_data = direct_response.json()
                        if direct_data.get('msg') == 'OK':
                            return direct_data['result']['url']
                except:
                    pass
                        
                # 如果無法獲取直接連結，返回網頁連結
                return f"https://katfile.cloud/{file_code}"
                
            except Exception as error:
                error_msg = f"❌ 上傳錯誤 (嘗試 {attempt + 1}/{max_retries}): {str(error)}"
                self.root.after(0, lambda msg=error_msg: self.log(msg))
                if attempt == max_retries - 1:
                    return None
                    
        return None
    
    def update_file_status(self, index, status):
        """更新檔案狀態顯示"""
        items = self.file_tree.get_children()
        if index < len(items):
            item = items[index]
            current_values = list(self.file_tree.item(item, "values"))
            current_values[1] = status
            self.file_tree.item(item, values=current_values)

def main():
    root = tk.Tk()
    app = KatFileUploaderEnhanced(root)
    root.mainloop()

if __name__ == "__main__":
    main()
